
doc=open("ab.txt",'r')
a=doc.readlines()
dc=open("docs.txt",'r')
b=dc.readlines()


import docx

print(a[3]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[0])
print(a[3+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[1])
print(a[4+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[2])
print(a[5+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[3])
print(a[6+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[4])
print(a[9+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[5])
print(a[10+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[6])
print(a[11+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[7])
print(a[20+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[8])
print(a[21+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[9])
print(a[22+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[10])
print(a[23+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[11])
print(a[24+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[12])
print(a[25+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[13])

myfile=docx.Document()

myfile.add_paragraph
myfile.add_paragraph(a[3]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[0])
myfile.add_paragraph(a[3+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[1])
myfile.add_paragraph(a[4+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[2])
myfile.add_paragraph(a[5+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[3])
myfile.add_paragraph(a[6+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[4])
myfile.add_paragraph(a[9+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[5])
myfile.add_paragraph(a[10+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[6])
myfile.add_paragraph(a[11+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[7])
myfile.add_paragraph(a[20+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[8])
myfile.add_paragraph(a[21+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[9])
myfile.add_paragraph(a[22+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[10])
myfile.add_paragraph(a[23+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[11])
myfile.add_paragraph(a[24+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[12])
myfile.add_paragraph(a[25+1]+"\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t\t="+b[13])

myfile.save("calculations.docx")